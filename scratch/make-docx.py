import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

doc = docx.Document()

# Page setup - Margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Color Palette: Corporate Navy / Steel / Emerald
COLOR_PRIMARY = RGBColor(26, 54, 93)      # Navy #1A365D
COLOR_SECONDARY = RGBColor(43, 108, 176)  # Blue #2B6CB0
COLOR_DARK = RGBColor(45, 55, 72)         # Charcoal #2D3748
COLOR_MUTED = RGBColor(113, 128, 150)     # Gray #718096
HEX_PRIMARY = "1A365D"
HEX_LIGHT_BG = "F7FAFC"
HEX_BORDER = "CBD5E0"
HEX_HEADER_BG = "2B6CB0"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

# Document Title
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(4)
r_org = title_p.add_run("COAMA SudAmérica — LUMO Data Solutions\n")
r_org.font.size = Pt(11)
r_org.font.bold = True
r_org.font.color.rgb = COLOR_SECONDARY

r_title = title_p.add_run("Especificación Técnica de Exportación a ERP (CSV)")
r_title.font.size = Pt(20)
r_title.font.bold = True
r_title.font.color.rgb = COLOR_PRIMARY

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(18)
r_sub = sub_p.add_run("Sistema de Registro de Tiempos Muertos y Paradas de Secaderos")
r_sub.font.size = Pt(12)
r_sub.font.italic = True
r_sub.font.color.rgb = COLOR_MUTED

# Divider
div_table = doc.add_table(rows=1, cols=1)
div_table.alignment = WD_TABLE_ALIGNMENT.CENTER
div_cell = div_table.cell(0, 0)
div_cell.width = Inches(6.9)
set_cell_background(div_cell, HEX_PRIMARY)
set_cell_margins(div_cell, top=20, bottom=20, left=0, right=0)
p_empty = div_cell.paragraphs[0]
p_empty.paragraph_format.space_before = Pt(0)
p_empty.paragraph_format.space_after = Pt(0)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 1: Resumen y Parámetros
h1 = doc.add_heading("1. Parámetros Globales del Archivo", level=1)
h1.runs[0].font.color.rgb = COLOR_PRIMARY
h1.runs[0].font.size = Pt(14)
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(6)

p_desc = doc.add_paragraph(
    "Este documento establece el estándar definitivo para la generación del archivo CSV de exportación "
    "de paradas operativas de los Secaderos (OMECO, BENEKE, RAUTE). Su diseño garantiza compatibilidad nativa "
    "con Microsoft Excel en español (región Argentina / Latinoamérica) y los motores de importación del sistema ERP."
)
p_desc.paragraph_format.space_after = Pt(8)

param_table = doc.add_table(rows=6, cols=3)
param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Parámetro", "Valor Especificado", "Descripción / Impacto Técnico"]
widths = [Inches(1.8), Inches(2.2), Inches(2.9)]

for j, h_text in enumerate(headers):
    cell = param_table.cell(0, j)
    cell.width = widths[j]
    set_cell_background(cell, HEX_PRIMARY)
    set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(h_text)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(9.5)

param_rows = [
    ("Separador de campos", "Punto y coma (;)", "Evita colisión con la coma decimal utilizada en Excel regional."),
    ("Codificación", "UTF-8 con BOM (\\ufeff)", "Fuerza a Excel Windows a detectar acentos (Ó, É, Ñ) sin corromper caracteres."),
    ("Final de línea", "Windows CRLF (\\r\\n)", "Estándar requerido para compatibilidad con servidores Windows."),
    ("Nombres de archivo", "COAMA_Exportacion_ERP.csv\nLUMO_SECADEROS_PARADAS_YYYY-MM-DD.csv", "Nombre fijo oficial y nombre versionado por fecha diaria."),
    ("Ubicación de salida", "output to erp/", "Carpeta local accesible en la raíz del proyecto para tareas automáticas."),
]

for i, row in enumerate(param_rows, start=1):
    bg_color = HEX_LIGHT_BG if i % 2 == 1 else "FFFFFF"
    for j, val in enumerate(row):
        cell = param_table.cell(i, j)
        cell.width = widths[j]
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(val)
        r.font.size = Pt(9)
        if j == 0:
            r.font.bold = True
            r.font.color.rgb = COLOR_DARK

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 2: Reglas de Tratamiento
h2 = doc.add_heading("2. Reglas de Tratamiento de Datos y Formato", level=1)
h2.runs[0].font.color.rgb = COLOR_PRIMARY
h2.runs[0].font.size = Pt(14)
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after = Pt(6)

rules = [
    ("Punto y coma (;) en textos:", "Cualquier ';' escrito por el operario en observaciones se reemplaza automáticamente por una coma (',') para evitar fragmentar columnas."),
    ("Saltos de línea en textos:", "Todo salto de línea (\\r, \\n) se convierte en espacio simple (' ') para preservar la fila física."),
    ("Comillas dobles (\"):", "Se escapan duplicándolas (\"\") y el campo se entrecomilla conforme al estándar RFC 4180."),
    ("Valores nulos / vacíos:", "Si no hay observaciones se exporta explícitamente '-.-'. Si la ubicación no aplica se deja vacía (';;')."),
    ("Separador decimal con coma (,):", "Todos los campos numéricos con decimales usan coma (,). Ejemplos: 0,01 horas; 0,4 minutos; 2,1 minutos."),
    ("Enteros sin ceros superfluos:", "Los valores enteros se emiten limpios (12 para horas programadas de turno; 0 cuando una duración redondea a cero)."),
    ("Deduplicación estricta de paradas:", "Se filtran los registros técnicos de fin (inicio_evento_id). Cada parada cerrada se exporta exactamente una única vez (sin duplicados)."),
]

for title, desc in rules:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title + " ")
    r1.font.bold = True
    r1.font.color.rgb = COLOR_DARK
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(9.5)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 3: Tabla de 13 Columnas
h3 = doc.add_heading("3. Tabla de Especificación de los 13 Campos", level=1)
h3.runs[0].font.color.rgb = COLOR_PRIMARY
h3.runs[0].font.size = Pt(14)
h3.paragraph_format.space_before = Pt(14)
h3.paragraph_format.space_after = Pt(6)

table = doc.add_table(rows=14, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
col_headers = ["#", "Col", "Nombre del Campo", "Tipo", "Formato", "Ejemplo"]
t_widths = [Inches(0.4), Inches(0.4), Inches(2.2), Inches(0.9), Inches(1.6), Inches(1.4)]

for j, h_text in enumerate(col_headers):
    cell = table.cell(0, j)
    cell.width = t_widths[j]
    set_cell_background(cell, HEX_HEADER_BG)
    set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(h_text)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(8.5)

fields_data = [
    ("1", "A", "fecha_de_registro", "Texto", "YYYY-MM-DD (10 car)", "2026-08-28"),
    ("2", "B", "linea", "Texto", "Alfanumérico (Mayúsc)", "OMECO"),
    ("3", "C", "turno_hora_desde", "Hora", "HH:MM:SS (8 car)", "06:00:00"),
    ("4", "D", "turno_hora_hasta", "Hora", "HH:MM:SS (8 car)", "18:00:00"),
    ("5", "E", "tiempo_de_turno_en_horas_programadas", "Numérico", "Entero / Decimal", "12"),
    ("6", "F", "categoria", "Texto", "Alfanumérico (Mayúsc)", "ELECTRICO"),
    ("7", "G", "tiempo_muerto", "Texto", "Alfanumérico (Mayúsc)", "CARGADOR"),
    ("8", "H", "observacion", "Texto", "Alfanumérico (-.- si vacío)", "-.-"),
    ("9", "I", "ubicacion", "Texto", "Alfanumérico (o vacío)", "N1P1"),
    ("10", "J", "tiempo_muerto_hora_desde", "ISO", "YYYY-MM-DDTHH:MM:SS.mmmZ", "2026-08-28T14:57:27.729Z"),
    ("11", "K", "tiempo_muerto_hora_hasta", "ISO", "YYYY-MM-DDTHH:MM:SS.mmmZ", "2026-08-28T14:57:51.591Z"),
    ("12", "L", "tiempo_muerto_en_horas", "Numérico", "2 decimales con coma (o 0)", "0,01"),
    ("13", "M", "tiempo_muerto_en_minutos", "Numérico", "1 decimal con coma (o 0)", "0,4"),
]

for i, row in enumerate(fields_data, start=1):
    bg_color = HEX_LIGHT_BG if i % 2 == 1 else "FFFFFF"
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.width = t_widths[j]
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(val)
        r.font.size = Pt(8.5)
        if j == 2:
            r.font.bold = True
            r.font.color.rgb = COLOR_PRIMARY

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 4: Muestra Real
h4 = doc.add_heading("4. Muestra del Archivo CSV Generado", level=1)
h4.runs[0].font.color.rgb = COLOR_PRIMARY
h4.runs[0].font.size = Pt(14)
h4.paragraph_format.space_before = Pt(14)
h4.paragraph_format.space_after = Pt(6)

sample_csv = (
    "fecha_de_registro;linea;turno_hora_desde;turno_hora_hasta;tiempo_de_turno_en_horas_programadas;categoria;tiempo_muerto;observacion;ubicacion;tiempo_muerto_hora_desde;tiempo_muerto_hora_hasta;tiempo_muerto_en_horas;tiempo_muerto_en_minutos\n"
    "2026-08-28;OMECO;06:00:00;18:00:00;12;ELECTRICO;CARGADOR;-.-;;2026-08-28T14:57:27.729Z;2026-08-28T14:57:51.591Z;0,01;0,4\n"
    "2026-08-28;OMECO;06:00:00;18:00:00;12;MECANICO;CARGADOR;-.-;;2026-08-28T14:58:12.621Z;2026-08-28T14:58:18.271Z;0;0,1\n"
    "2026-08-28;OMECO;06:00:00;18:00:00;12;PROCESO;SECADERO TRANCADO;-.-;N1P1;2026-08-28T15:00:01.402Z;2026-08-28T15:00:21.105Z;0,01;0,3\n"
    "2026-08-28;OMECO;06:00:00;18:00:00;12;LOGISTICA;PRUEBA;-.-;;2026-08-28T15:01:55.904Z;2026-08-28T15:02:07.421Z;0;0,2\n"
    "2026-08-28;OMECO;06:00:00;18:00:00;12;MECANICO;CADENA;-.-;;2026-08-28T15:04:42.121Z;2026-08-28T15:05:03.001Z;0,01;0,3\n"
    "2026-08-28;OMECO;06:00:00;18:00:00;12;OPERATIVO;CALDERA;-.-;;2026-08-28T15:06:30.764Z;2026-08-28T15:06:38.981Z;0;0,1\n"
    "2026-08-28;OMECO;06:00:00;18:00:00;12;PROCESO;SECADERO TRANCADO;-.-;N6P2;2026-08-28T15:09:07.003Z;2026-08-28T15:09:42.588Z;0,01;0,6"
)

code_table = doc.add_table(rows=1, cols=1)
code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
c_cell = code_table.cell(0, 0)
c_cell.width = Inches(6.9)
set_cell_background(c_cell, HEX_LIGHT_BG)
set_cell_margins(c_cell, top=120, bottom=120, left=140, right=140)
p_code = c_cell.paragraphs[0]
p_code.paragraph_format.space_before = Pt(0)
p_code.paragraph_format.space_after = Pt(0)
r_code = p_code.add_run(sample_csv)
r_code.font.name = "Consolas"
r_code.font.size = Pt(7.5)
r_code.font.color.rgb = COLOR_DARK

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 5: Métodos de Obtención
h5 = doc.add_heading("5. Métodos de Disponibilidad y Consumo", level=1)
h5.runs[0].font.color.rgb = COLOR_PRIMARY
h5.runs[0].font.size = Pt(14)
h5.paragraph_format.space_before = Pt(14)
h5.paragraph_format.space_after = Pt(6)

methods = [
    ("1. Archivo por lotes (Escritorio / Servidor):", "Doble clic en 'Exportar_Para_ERP.bat' o ejecutando 'npm run export:erp'. Genera el archivo en la carpeta 'output to erp/'."),
    ("2. Descarga desde el Portal Supervisor:", "Haciendo clic en el botón 'Descargar CSV ERP' en la pestaña de Análisis del portal web."),
    ("3. Endpoint HTTP (Automatización para el ERP):", "GET http://<IP_SERVIDOR>:8080/export/erp — Permite que tareas programadas en el servidor ERP (PowerShell, cURL, Python) descarguen el archivo directamente."),
]

for title, desc in methods:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + " ")
    r1.font.bold = True
    r1.font.color.rgb = COLOR_DARK
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(9.5)

# Output destinations
out_paths = [
    r"e:\Jorge Falcon dbd\COAMA\Desarrollo\APP SECADEROS - COPIA\coama-secaderos-2\docs\technical\COAMA_Especificacion_Exportacion_ERP.docx",
    r"e:\Jorge Falcon dbd\COAMA\Desarrollo\APP SECADEROS - COPIA\docs\technical\COAMA_Especificacion_Exportacion_ERP.docx",
    r"e:\Jorge Falcon dbd\COAMA\Desarrollo\APP SECADEROS - COPIA\coama-secaderos-2\output to erp\COAMA_Especificacion_Exportacion_ERP.docx",
]

for out_path in out_paths:
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"[OK] Guardado: {out_path}")
